"""
Multimodal parsers. Each parser takes a file path and returns a list of
`ParsedChunk` objects — raw, un-scrubbed text plus metadata. Scrubbing
happens as a separate pipeline stage (see pipeline.py) so parsers stay
single-purpose and testable in isolation.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

Modality = Literal["text", "table", "image_ocr", "audio_transcript"]


@dataclass
class ParsedChunk:
    text: str
    modality: Modality
    source_path: str
    page_or_segment: int | None = None
    metadata: dict = field(default_factory=dict)


def parse_pdf(path: str | Path) -> list[ParsedChunk]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    chunks = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(
                ParsedChunk(text=text, modality="text", source_path=str(path), page_or_segment=i)
            )
    return chunks


def parse_docx(path: str | Path) -> list[ParsedChunk]:
    import docx

    doc = docx.Document(str(path))
    chunks = []

    # paragraphs
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if full_text.strip():
        chunks.append(ParsedChunk(text=full_text, modality="text", source_path=str(path)))

    # tables -> flatten to a readable text block per table
    for t_idx, table in enumerate(doc.tables):
        rows = ["\t".join(cell.text for cell in row.cells) for row in table.rows]
        table_text = "\n".join(rows)
        if table_text.strip():
            chunks.append(
                ParsedChunk(
                    text=table_text,
                    modality="table",
                    source_path=str(path),
                    page_or_segment=t_idx,
                    metadata={"table_index": t_idx},
                )
            )
    return chunks


def parse_image(path: str | Path) -> list[ParsedChunk]:
    """OCR-based text extraction from images (screenshots, scanned docs, etc.)"""
    import pytesseract
    from PIL import Image

    img = Image.open(str(path))
    text = pytesseract.image_to_string(img)
    if not text.strip():
        return []
    return [ParsedChunk(text=text, modality="image_ocr", source_path=str(path))]


def parse_audio(path: str | Path, model_size: str = "base") -> list[ParsedChunk]:
    """
    Whisper transcription. `model_size="base"` is a reasonable CPU default;
    "tiny" is faster/lower quality, "small"/"medium" are slower and much
    better. Tune per your hardware in slm/config.py conventions.
    """
    import whisper

    model = whisper.load_model(model_size)
    result = model.transcribe(str(path))

    chunks = []
    for seg in result.get("segments", []):
        text = seg.get("text", "").strip()
        if text:
            chunks.append(
                ParsedChunk(
                    text=text,
                    modality="audio_transcript",
                    source_path=str(path),
                    page_or_segment=seg.get("id"),
                    metadata={"start": seg.get("start"), "end": seg.get("end")},
                )
            )
    return chunks


PARSER_BY_EXTENSION = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".png": parse_image,
    ".jpg": parse_image,
    ".jpeg": parse_image,
    ".mp3": parse_audio,
    ".wav": parse_audio,
    ".m4a": parse_audio,
}


def parse_file(path: str | Path) -> list[ParsedChunk]:
    path = Path(path)
    ext = path.suffix.lower()
    parser = PARSER_BY_EXTENSION.get(ext)
    if parser is None:
        raise ValueError(f"No parser registered for extension '{ext}' ({path})")
    return parser(path)
